import os
import PyPDF2
import docx

class ResumeProcessor:
    """Service for processing and extracting information from resumes"""
    
    def __init__(self):
        """Initialize the resume processor"""
        pass
    
    def extract_text(self, resume_path):
        """
        Extract text from resume file (PDF, DOCX, or TXT)
        
        Args:
            resume_path (str): Path to the resume file
            
        Returns:
            str: Extracted text from the resume
        """
        file_extension = os.path.splitext(resume_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(resume_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_from_docx(resume_path)
        elif file_extension == '.txt':
            return self._extract_from_txt(resume_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, pdf_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Check if text extraction was successful
                        text += page_text + "\n"
                    else:
                        text += "[PDF page with no extractable text]\n"
        except Exception as e:
            text = f"Error extracting PDF text: {str(e)}"
        return text
    
    def _extract_from_docx(self, docx_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            text = f"Error extracting DOCX text: {str(e)}"
        return text
    
    def _extract_from_txt(self, txt_path):
        """Extract text from TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            return f"Error extracting TXT text: {str(e)}"
    
    def extract_sections(self, resume_text):
        """
        Attempt to extract common resume sections
        
        Args:
            resume_text (str): The full resume text
            
        Returns:
            dict: Dictionary of resume sections
        """
        sections = {
            'contact_info': '',
            'education': '',
            'experience': '',
            'skills': '',
            'projects': '',
            'certifications': '',
            'other': ''
        }
        
        # Simple heuristic-based section extraction
        lines = resume_text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line = line.strip()
            lower_line = line.lower()
            
            # Check for section headers
            if any(keyword in lower_line for keyword in ['education', 'academic', 'degree']):
                if len(line) < 100:  # Avoid treating paragraphs as headers
                    current_section = 'education'
                    continue
            elif any(keyword in lower_line for keyword in ['experience', 'employment', 'work history']):
                if len(line) < 100:
                    current_section = 'experience'
                    continue
            elif any(keyword in lower_line for keyword in ['skill', 'technical', 'technologies', 'languages']):
                if len(line) < 100:
                    current_section = 'skills'
                    continue
            elif any(keyword in lower_line for keyword in ['project', 'portfolio']):
                if len(line) < 100:
                    current_section = 'projects'
                    continue
            elif any(keyword in lower_line for keyword in ['certification', 'certificate', 'license']):
                if len(line) < 100:
                    current_section = 'certifications'
                    continue
            elif any(keyword in lower_line for keyword in ['contact', 'email', 'phone', 'address']):
                if len(line) < 100:
                    current_section = 'contact_info'
                    continue
            
            # Add line to current section
            if line:
                sections[current_section] += line + '\n'
        
        return sections
